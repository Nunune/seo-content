import re, sys, os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image as PILImage

FONT_NAME = 'Arial'
FONT_SIZE_BODY = 11
FONT_SIZE_H1 = 18
FONT_SIZE_H2 = 15
FONT_SIZE_H3 = 13
COLOR_H1 = RGBColor(0x1D, 0x4E, 0xD8)
COLOR_H2 = RGBColor(0x1E, 0x40, 0xAF)
COLOR_H3 = RGBColor(0x37, 0x41, 0x51)
MAX_IMG_WIDTH = Inches(5.5)

preprocessed_path = sys.argv[1]
output_path = sys.argv[2]

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(FONT_SIZE_BODY)
style.paragraph_format.space_after = Pt(6)

def set_heading_style(para, level):
    sizes = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3}
    colors = {1: COLOR_H1, 2: COLOR_H2, 3: COLOR_H3}
    for run in para.runs:
        run.font.size = Pt(sizes.get(level, 12))
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
        run.font.bold = True
        run.font.name = FONT_NAME

def set_img_alt(inline, alt_text):
    try:
        pic = inline._inline.findall('.//' + qn('pic:pic'))[0]
        nvPicPr = pic.find(qn('pic:nvPicPr'))
        cNvPr = nvPicPr.find(qn('pic:cNvPr'))
        cNvPr.set('descr', alt_text)
    except Exception:
        pass

def _add_runs(para, text, bold=False, italic=False, color=None):
    LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    pos = 0
    for m in LINK.finditer(text):
        before = text[pos:m.start()]
        if before:
            run = para.add_run(before)
            run.font.name = FONT_NAME
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color
        run = para.add_run(m.group(1))
        run.font.name = FONT_NAME
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        pos = m.end()
    tail = text[pos:]
    if tail:
        run = para.add_run(tail)
        run.font.name = FONT_NAME
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

def parse_inline(para, text, color=None):
    TOKEN = re.compile(
        r'\*\*(.+?)\*\*'
        r'|\[([^\]]+)\]\(([^)]+)\)'
        r'|\*([^*\n]+)\*',
        re.DOTALL
    )
    pos = 0
    for m in TOKEN.finditer(text):
        before = text[pos:m.start()]
        if before:
            _add_runs(para, before, color=color)
        if m.group(1) is not None:
            _add_runs(para, m.group(1), bold=True, color=color)
        elif m.group(2) is not None:
            run = para.add_run(m.group(2))
            run.font.name = FONT_NAME
            run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        elif m.group(4) is not None:
            _add_runs(para, m.group(4), italic=True, color=color)
        pos = m.end()
    tail = text[pos:]
    if tail:
        _add_runs(para, tail, color=color)

def add_paragraph(text, style_name='Normal'):
    p = doc.add_paragraph(style=style_name)
    parse_inline(p, text)
    return p

with open(preprocessed_path, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
table_buffer = []

while i < len(lines):
    line = lines[i].rstrip('\n')
    stripped = line.strip()

    # Image placeholder
    if stripped.startswith('> 📷 **[CHỖ NÀY CHÈN ẢNH:'):
        p = doc.add_paragraph()
        run = p.add_run(stripped[2:].replace('**', ''))
        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
        run.italic = True
        run.font.name = FONT_NAME
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        i += 1
        continue

    # Markdown image
    img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
    if img_match:
        alt = img_match.group(1)
        path = img_match.group(2)
        if os.path.exists(path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                if path.lower().endswith('.webp'):
                    with PILImage.open(path) as pil_img:
                        if pil_img.mode in ('RGBA', 'LA', 'P'):
                            pil_img = pil_img.convert('RGB')
                        buf = BytesIO()
                        pil_img.save(buf, format='JPEG', quality=90)
                        buf.seek(0)
                    inline = run.add_picture(buf, width=MAX_IMG_WIDTH)
                else:
                    inline = run.add_picture(path, width=MAX_IMG_WIDTH)
                set_img_alt(inline, alt)
            except Exception as e:
                doc.add_paragraph(f'[Không tải được ảnh: {path} — {e}]')
        i += 1
        continue

    # Headings
    if stripped.startswith('### '):
        p = doc.add_heading(stripped[4:], level=3)
        set_heading_style(p, 3)
        i += 1
        continue
    if stripped.startswith('## '):
        p = doc.add_heading(stripped[3:], level=2)
        set_heading_style(p, 2)
        i += 1
        continue
    if stripped.startswith('# '):
        p = doc.add_heading(stripped[2:], level=1)
        set_heading_style(p, 1)
        i += 1
        continue

    # Horizontal rule
    if stripped in ('---', '***', '___'):
        p = doc.add_paragraph('─' * 60)
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
        i += 1
        continue

    # Blockquote bullet: > - text
    if stripped.startswith('> - ') or stripped.startswith('> * '):
        text = stripped[4:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        bullet_run = p.add_run('• ')
        bullet_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        bullet_run.font.name = FONT_NAME
        bullet_run.font.size = Pt(10)
        parse_inline(p, text)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        i += 1
        continue

    # Blockquote plain: > text
    if stripped.startswith('> '):
        text = stripped[2:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        parse_inline(p, text)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        i += 1
        continue

    # Table detection
    if '|' in stripped and stripped.startswith('|'):
        table_buffer = [stripped]
        i += 1
        while i < len(lines):
            tline = lines[i].rstrip('\n').strip()
            if '|' in tline and tline.startswith('|'):
                table_buffer.append(tline)
                i += 1
            else:
                break
        rows = []
        for trow in table_buffer:
            if re.match(r'^\|[-| :]+\|$', trow):
                continue
            cells = [c.strip() for c in trow.split('|') if c.strip()]
            if cells:
                rows.append(cells)
        if rows:
            col_count = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = 'Table Grid'
            for r_i, row_data in enumerate(rows):
                for c_i, cell_text in enumerate(row_data):
                    if c_i < col_count:
                        cell = table.rows[r_i].cells[c_i]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        parse_inline(p, cell_text)
                        for run in p.runs:
                            run.font.name = FONT_NAME
                            run.font.size = Pt(10)
                        if r_i == 0:
                            for run in p.runs:
                                run.bold = True
            doc.add_paragraph()
        continue

    # Bullet list
    if stripped.startswith('- ') or stripped.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        parse_inline(p, stripped[2:])
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = Pt(FONT_SIZE_BODY)
        i += 1
        continue

    # Numbered list
    num_match = re.match(r'^(\d+)\. (.+)', stripped)
    if num_match:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        num_run = p.add_run(num_match.group(1) + '. ')
        num_run.font.name = FONT_NAME
        num_run.font.size = Pt(FONT_SIZE_BODY)
        parse_inline(p, num_match.group(2))
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = Pt(FONT_SIZE_BODY)
        i += 1
        continue

    # Regular paragraph
    if stripped:
        p = add_paragraph(stripped)
        p.paragraph_format.space_after = Pt(8)

    i += 1

doc.save(output_path)
print('Saved:', output_path)
